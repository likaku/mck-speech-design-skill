#!/usr/bin/env python3
"""
speech_to_docx.py — Convert a speech markdown file to a formatted Word document.

Usage:
    python speech_to_docx.py <speech.md> [output.docx]

Arguments:
    speech.md   : Path to the speech markdown file
    output.docx : (Optional) Output path. Defaults to <speech_stem>.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ──────────────────────────────────────────────────────────
# Style constants
# ──────────────────────────────────────────────────────────

FONT_BODY = '微软雅黑'
FONT_BODY_FALLBACK = 'Microsoft YaHei'
FONT_HEADING = '微软雅黑'
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x2E)   # Dark navy
COLOR_ACCENT = RGBColor(0x00, 0x6B, 0xFF)     # Accent blue
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_TABLE_HEADER = RGBColor(0xF0, 0xF4, 0xFF)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def _add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """Add a run to a paragraph with formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def _configure_styles(doc):
    """Configure document styles for the speech document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(11)
    font.color.rgb = COLOR_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35

    # Headings
    for level, size in [(1, 22), (2, 16), (3, 13)]:
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_HEADING
        h.font.size = Pt(size)
        h.font.color.rgb = COLOR_HEADING
        h.font.bold = True
        h.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(8)


# ──────────────────────────────────────────────────────────
# Markdown parser (lightweight, speech-specific)
# ──────────────────────────────────────────────────────────

def parse_speech_md(md_text: str) -> dict:
    """Parse the speech markdown into structured sections."""
    result = {
        'title': '',
        'meta': [],         # list of "key: value" strings from blockquote header
        'sections': [],     # list of {heading, level, content_blocks}
        'raw_lines': md_text.split('\n'),
    }

    lines = md_text.split('\n')
    i = 0

    # Title (first # heading)
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# ') and not line.startswith('## '):
            result['title'] = line[2:].strip()
            i += 1
            break
        i += 1

    # Meta block (blockquote lines right after title)
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('>'):
            meta_text = line.lstrip('> ').strip()
            if meta_text:
                result['meta'].append(meta_text)
            i += 1
        elif line == '':
            i += 1
        else:
            break

    # Sections (## and ### headings with content)
    current_section = None
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('## '):
            if current_section:
                result['sections'].append(current_section)
            current_section = {
                'heading': stripped[3:].strip(),
                'level': 2,
                'blocks': [],
            }
        elif stripped.startswith('### '):
            if current_section:
                current_section['blocks'].append({
                    'type': 'subheading',
                    'text': stripped[4:].strip(),
                })
        elif stripped.startswith('|') and current_section:
            # Table row
            if not current_section['blocks'] or current_section['blocks'][-1]['type'] != 'table':
                current_section['blocks'].append({'type': 'table', 'rows': []})
            # Skip separator rows (|---|---|)
            if not re.match(r'^\|[\s\-:|]+\|$', stripped):
                cells = [c.strip() for c in stripped.strip('|').split('|')]
                current_section['blocks'][-1]['rows'].append(cells)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            if current_section:
                current_section['blocks'].append({
                    'type': 'bullet',
                    'text': stripped[2:].strip(),
                })
        elif stripped.startswith('> '):
            if current_section:
                current_section['blocks'].append({
                    'type': 'quote',
                    'text': stripped[2:].strip(),
                })
        elif stripped == '---':
            pass  # Skip horizontal rules
        elif stripped and current_section:
            # Regular paragraph text
            if (current_section['blocks']
                    and current_section['blocks'][-1]['type'] == 'paragraph'):
                current_section['blocks'][-1]['text'] += '\n' + stripped
            else:
                current_section['blocks'].append({
                    'type': 'paragraph',
                    'text': stripped,
                })
        elif stripped == '' and current_section and current_section['blocks']:
            # Blank line — end current paragraph
            last = current_section['blocks'][-1]
            if last['type'] == 'paragraph':
                pass  # Next non-blank will create a new paragraph

        i += 1

    if current_section:
        result['sections'].append(current_section)

    return result


def _clean_md_formatting(text: str) -> str:
    """Strip markdown bold/italic markers for plain docx text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = text.replace('`[pause]`', '').replace('[pause]', '')
    return text.strip()


# ──────────────────────────────────────────────────────────
# Document builder
# ──────────────────────────────────────────────────────────

def build_docx(parsed: dict, output_path: Path):
    """Build a formatted Word document from parsed speech data."""
    doc = Document()
    _configure_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    _add_run(title_para, parsed['title'], bold=True, size=24, color=COLOR_HEADING, font_name=FONT_HEADING)

    # ── Meta info ──
    if parsed['meta']:
        for meta_line in parsed['meta']:
            meta_clean = _clean_md_formatting(meta_line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, meta_clean, size=10, color=COLOR_MUTED, font_name=FONT_BODY)
            p.paragraph_format.space_after = Pt(2)

    # Separator
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sep, '─' * 50, color=RGBColor(0xCC, 0xCC, 0xCC), size=8)

    # ── Sections ──
    for section_data in parsed['sections']:
        heading_text = section_data['heading']
        # Use Heading 1 for ## level
        h = doc.add_heading(heading_text, level=1 if section_data['level'] == 2 else 2)

        for block in section_data['blocks']:
            btype = block['type']

            if btype == 'subheading':
                doc.add_heading(block['text'], level=2)

            elif btype == 'table':
                rows_data = block['rows']
                if not rows_data:
                    continue
                ncols = len(rows_data[0])
                table = doc.add_table(rows=0, cols=ncols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ri, row_cells in enumerate(rows_data):
                    row = table.add_row()
                    for ci, cell_text in enumerate(row_cells):
                        cell = row.cells[ci]
                        cell.text = _clean_md_formatting(cell_text)
                        # Style cell text
                        for para in cell.paragraphs:
                            para.paragraph_format.space_after = Pt(2)
                            for run in para.runs:
                                run.font.size = Pt(10)
                                run.font.name = FONT_BODY
                                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
                        # Header row shading
                        if ri == 0:
                            _set_cell_shading(cell, 'F0F4FF')
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True

                doc.add_paragraph()  # spacing after table

            elif btype == 'bullet':
                clean = _clean_md_formatting(block['text'])
                p = doc.add_paragraph(clean, style='List Bullet')
                for run in p.runs:
                    run.font.name = FONT_BODY
                    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

            elif btype == 'quote':
                clean = _clean_md_formatting(block['text'])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                _add_run(p, clean, italic=True, color=COLOR_ACCENT, size=11, font_name=FONT_BODY)

            elif btype == 'paragraph':
                clean = _clean_md_formatting(block['text'])
                if not clean:
                    continue
                p = doc.add_paragraph()
                # Check for Q&A format
                if clean.startswith('Q:') or clean.startswith('**Q:'):
                    _add_run(p, clean, bold=True, size=11, font_name=FONT_BODY)
                elif clean.startswith('A:') or clean.startswith('**A:'):
                    _add_run(p, clean, size=11, color=COLOR_MUTED, font_name=FONT_BODY)
                else:
                    _add_run(p, clean, size=11, font_name=FONT_BODY)

    # ── Footer ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(30)
    _add_run(footer_para, '— END —', size=10, color=COLOR_MUTED, font_name=FONT_BODY)

    doc.save(str(output_path))
    print(f"✅ Word document saved: {output_path}")
    return str(output_path)


# ──────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out_path = Path(sys.argv[2])
    else:
        out_path = md_path.parent / f'{md_path.stem}.docx'

    md_text = md_path.read_text(encoding='utf-8')
    parsed = parse_speech_md(md_text)
    build_docx(parsed, out_path)


if __name__ == '__main__':
    main()
